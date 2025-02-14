import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import tempfile
from num2words import num2words
import uuid 

# Common Functions
def apply_formatting(run, font_name, font_size, bold=False):
    """Apply specific formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def replace_and_format(doc, placeholders, font_name, font_size, option):
    """Replace placeholders and apply formatting."""
    for para in doc.paragraphs:
        if para.text:
            for key, value in placeholders.items():
                if key in para.text:
                    runs = para.runs
                    for run in runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            if para == doc.paragraphs[0]:
                                apply_formatting(run, font_name, font_size, bold=True)
                        else:
                            run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    for key, value in placeholders.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if key == "<<Address>>" else WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    apply_formatting(run, "Times New Roman", 11 if option == "NDA" else 12)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def edit_word_template(template_path, output_path, placeholders, font_name, font_size, option):
    """Edit Word document and apply formatting."""
    try:
        doc = Document(template_path)
        replace_and_format(doc, placeholders, font_name, font_size, option)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")

def apply_image_placeholder(doc, placeholder_key, image_file):
    """Replace a placeholder with an image in the Word document."""
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder_key in para.text:
                            para.text = ""
                            run = para.add_run()
                            run.add_picture(image_file, width=Inches(1.5), height=Inches(0.75))
                            return doc
        for para in doc.paragraphs:
            if placeholder_key in para.text:
                para.text = ""
                run = para.add_run()
                run.add_picture(image_file, width=Inches(1.2), height=Inches(0.10))
                return doc
        raise ValueError(f"Placeholder '{placeholder_key}' not found in the document.")
    except Exception as e:
        raise Exception(f"Error inserting image: {e}")

# Contract/NDA Generator
from docx.shared import Inches  # Add this import at the top of your script

def generate_document(option):
    """Streamlit UI for generating NDA or Contract documents."""
    st.title("Document Generator")

    base_dir = os.path.abspath(os.path.dirname(__file__))
    template_paths = {
        "NDA": os.path.join(base_dir, "NDA Template 1.docx"),
        "Contract": os.path.join(base_dir, "Contract Template 1.docx"),
    }

    client_name = st.text_input("Enter Client Name:")
    company_name = st.text_input("Enter Company Name:")
    address = st.text_area("Enter Address:")
    date_field = st.date_input("Enter Date:", datetime.today())

    # Add an input field for uploading an e-signature image
    signature_file = st.file_uploader("Upload E-Signature (PNG or JPEG)", type=["png", "jpg", "jpeg"])

    placeholders = {
        "<< Client Name >>": client_name,
        "<<Company Name>>": company_name,
        "<<Address>>": address,
        "<< Date >>": date_field.strftime("%d-%m-%Y"),
        "<< Date (Signature) >>": date_field.strftime("%d-%m-%Y"),
    }

    if st.button(f"Generate {option}") or st.session_state.get('buttons_visible', False):
        st.session_state.buttons_visible = True
        formatted_date = date_field.strftime("%d %b %Y")

        # Generate a unique file name using UUID
        unique_id = str(uuid.uuid4())[:8]  # Shorten UUID for readability
        file_name = f"{option} - {client_name} {formatted_date} - {unique_id}.docx"

        # Use a temporary directory
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, file_name)

        try:
            font_size = 11 if option == "NDA" else 12
            doc = Document(template_paths[option])

            # Replace placeholders in the document
            replace_and_format(doc, placeholders, "Times New Roman", font_size, option)

            # Replace <<Signature>> placeholder with the uploaded e-signature image
            if signature_file:
                # Save the uploaded signature to a temporary file
                signature_path = os.path.join(temp_dir, "signature.png")
                with open(signature_path, "wb") as f:
                    f.write(signature_file.getbuffer())

                # Use the apply_image_placeholder function to insert the signature
                doc = apply_image_placeholder(doc, "<<Signature>>", signature_path)

            # Save the updated document
            doc.save(output_path)
            st.success(f"{option} Document Generated Successfully!")

            # Provide download link for the generated document
            with open(output_path, "rb") as file:
                st.download_button(
                    label="Download Document (Word)",
                    data=file,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word"
                )

        except Exception as e:
            st.error(f"An error occurred: {e}")


# Invoice Generator
def format_price(amount, currency):
    """Format price based on currency."""
    formatted_price = f"{amount:,.2f}"
    return f"{currency} {formatted_price}" if currency == "USD" else f"Rs. {formatted_price}"

def format_percentage(value):
    """Format percentage without decimals."""
    return f"{int(value)}%"

def get_next_invoice_number():
    """Fetch and increment the invoice number."""
    invoice_file = "invoice_counter.txt"
    if not os.path.exists(invoice_file):
        with open(invoice_file, "w") as file:
            file.write("1000")  # Starting invoice number
    try:
        with open(invoice_file, "r") as file:
            current_invoice = int(file.read().strip())
    except ValueError:
        current_invoice = 1000
    next_invoice = current_invoice + 1
    with open(invoice_file, "w") as file:
        file.write(str(next_invoice))
    return current_invoice

def amount_to_words(amount):
    """Convert amount to words without currency formatting."""
    words = num2words(amount, lang='en').replace(',', '').title()
    return words

def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        # Apply bold formatting for specific placeholders
                        if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to Word>>":
                            run.bold = True  # Apply bold formatting

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    # Apply bold formatting for specific placeholders
                                    if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to Word>>":
                                        run.bold = True  # Apply bold formatting
    return doc

def edit_invoice_template(template_name, output_path, placeholders):
    """Edit an invoice template and save the result."""
    try:
        doc = Document(template_name)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing invoice template: {e}")

def generate_invoice():
    """Retrieve and update the next invoice number dynamically."""
    counter_file = "invoice_counter.txt"

    # Check if the file exists; if not, initialize it
    if not os.path.exists(counter_file):
        with open(counter_file, "w") as file:
            file.write("1")

    # Read the last invoice number
    with open(counter_file, "r") as file:
        last_number = int(file.read().strip())

    # Increment and update the file
    next_number = last_number + 1
    with open(counter_file, "w") as file:
        file.write(str(next_number))

    # Format invoice number (e.g., 001, 002, 003)
    return f"{next_number:03d}"

# Function to format prices based on the region
def format_price(amount, region):
    """Format price based on the region (INR or USD)."""
    return f"₹{amount:.2f}" if region == "INR" else f"${amount:.2f}"

# Function to convert numbers to words (placeholder function)
def amount_to_words(amount):
    """Convert amount to words (can integrate a library like num2words)."""
    return f"{amount} Rupees" if amount < 100000 else f"{amount} Dollars"

# Function to edit the invoice template
from docx import Document

def edit_invoice_template(template_path, output_path, placeholders):
    """Replace placeholders in the Word document template."""
    try:
        doc = Document(template_path)

        # Replace text in paragraphs
        for para in doc.paragraphs:
            for key, value in placeholders.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)

        # Replace text inside tables (for placeholders in table cells)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholders.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        # Save the modified document
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error processing invoice: {e}")
        return False


# Streamlit App
def generate_invoice():
    """Streamlit app for generating invoices."""
    st.title("Invoice Generator")

    # Select region
    region = st.selectbox("Select Region", ["INR", "USD"])

    # Set payment options dynamically
    payment_options = ["1 Payment", "3 EMI", "5 EMI"] if region == "INR" else ["3 EMI", "5 EMI"]

    # Input Fields
    client_name = st.text_input("Client Name")
    client_address = st.text_input("Client Address")
    project_name = st.text_input("Project Name")
    phone_number = st.text_input("Phone Number")
    gst_number = st.text_input("GST Number")
    base_amount = st.number_input("Base Amount (excluding GST)", min_value=0.0, format="%.2f")
    payment_option = st.selectbox("Payment Option", payment_options)
    invoice_date = st.date_input("Invoice Date", value=datetime.today())
    formatted_date = invoice_date.strftime("%d-%m-%Y")

    # Calculate GST and total amount
    gst_amount = round(base_amount * 0.18)
    total_amount = base_amount + gst_amount

    # Prepare placeholders for template
    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Address>>": client_address,
        "<<GST Number>>": gst_number,
        "<<Client Email>>": client_address,
        "<<Project Name>>": project_name,
        "<<Mobile Number>>": phone_number,
        "<<Date>>": formatted_date,
        "<<Amt to word>>": amount_to_words(int(total_amount)),
    }

    # Select the correct template based on payment option
    if payment_option == "1 Payment":
        template_name = f"Invoice Template - {region} - 1 Payment 1.docx"
        placeholders.update({
            "<<Price 1>>": format_price(base_amount, region),
            "<<Price 2>>": format_price(gst_amount, region),
            "<<Price 3>>": format_price(total_amount, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    elif payment_option == "3 EMI":
        template_name = f"Invoice Template - {region} - 3 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.30)
        p2 = round(total_amount * 0.40)
        p3 = total_amount - (p1 + p2)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(gst_amount, region),
            "<<Price 5>>": format_price(total_amount, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
        })

    elif payment_option == "5 EMI":
        template_name = f"Invoice Template - {region} - 5 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.20)
        p2 = round(total_amount * 0.20)
        p3 = round(total_amount * 0.20)
        p4 = round(total_amount * 0.20)
        p5 = total_amount - (p1 + p2 + p3 + p4)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(p4, region),
            "<<Price 5>>": format_price(p5, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
            "<<Price 9>>": format_price(p4, region),
            "<<Price 10>>": format_price(p5, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    # Generate Invoice Button
    if st.button("Generate Invoice"):
        try:
            # Generate the next invoice number
            invoice_number = get_next_invoice_number()
            placeholders["<<Invoice>>"] = str(invoice_number)
            placeholders["<<Invoice No>>"] = str(invoice_number)

            # Define the invoice template file path
            template_path = os.path.join(os.getcwd(), template_name)

            # Save the invoice to a temporary directory
            temp_dir = tempfile.gettempdir()
            sanitized_client_name = "".join([c if c.isalnum() or c.isspace() else "_" for c in client_name])
            output_path = os.path.join(temp_dir, f"Invoice_{sanitized_client_name}_{invoice_number}.docx")

            # Edit the template and save the invoice
            success = edit_invoice_template(template_path, output_path, placeholders)

            # Show success message and provide download link
            if success and os.path.exists(output_path):
                st.success(f"✅ Invoice #{invoice_number} generated successfully!")
                with open(output_path, "rb") as file:
                    st.download_button(
                        label="📥 Download Invoice",
                        data=file,
                        file_name=f"Invoice_{sanitized_client_name}_{invoice_number}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error(f"❌ Invoice generation failed.")
        except Exception as e:
            st.error(f"An error occurred: {e}")



# Main App
def main():
    st.sidebar.title("Select Application")
    app_choice = st.sidebar.radio("Choose an application:", ["Document Generator", "Invoice Generator"])

    if app_choice == "Document Generator":
        option = st.selectbox("Select Document Type", ["NDA", "Contract"], key="doc_type")
        generate_document(option)
    elif app_choice == "Invoice Generator":
        generate_invoice()

if __name__ == "__main__":
    main()
